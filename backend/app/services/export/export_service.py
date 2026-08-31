"""
Message export service.

Supports three formats:
  pdf   — reportlab PDF
  word  — python-docx .docx
  image — Pillow PNG screenshot of the answer text

All formats strip markdown syntax and render plain text.

ECharts charts are detected in answer text, rendered to PNG via pyecharts,
and embedded in the output (Word gets image + raw JSON, PDF/image get image only).
"""
from __future__ import annotations

import io
import json
import os
import re
import tempfile
import textwrap
from typing import Literal, Tuple

ExportFormat = Literal["pdf", "word", "image"]

# Regex to match echarts code blocks: ```echarts ... ```
_ECHARTS_BLOCK_RE = re.compile(r'```echarts\s*\n(.*?)\n```', re.DOTALL)


from app.services.infrastructure.reasoning_tags import strip_reasoning_tags

def _strip_markdown(text: str) -> str:
    """Very lightweight markdown stripper — enough for clean export."""
    text = strip_reasoning_tags(text)
    text = re.sub(r"#{1,6}\s*", "", text)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"`{1,3}(.*?)`{1,3}", r"\1", text, flags=re.DOTALL)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[citation:\d+\]", "", text)
    text = re.sub(r"\[citation\]\(\d+\)", "", text)
    return text.strip()


def _render_echarts_to_png(json_str: str) -> bytes | None:
    """Render an ECharts JSON config to PNG bytes using pyecharts.

    Returns None if rendering fails (malformed JSON, unsupported type, etc.).
    """
    try:
        config = json.loads(json_str)
    except (json.JSONDecodeError, ValueError):
        return None

    try:
        from pyecharts import options as opts
        from pyecharts.charts import (
            Bar, Line, Pie, Scatter, Radar, Gauge, Funnel,
            HeatMap, WordCloud, Boxplot, Candlestick, Parallel,
            Graph, Tree, Sankey, Chord, Liquid, Calendar,
            Polar, PictorialBar, ThemeRiver,
            Bar3D, Line3D, Scatter3D, Surface3D, Lines3D,
            Grid,
        )

        # Determine chart type from first series
        series_list = config.get("series", [])
        if not isinstance(series_list, list) or not series_list:
            return None

        chart_type = series_list[0].get("type", "bar")

        # Build a pyecharts chart instance based on type
        chart_map = {
            "bar": Bar,
            "line": Line,
            "pie": Pie,
            "scatter": Scatter,
            "radar": Radar,
            "gauge": Gauge,
            "funnel": Funnel,
            "heatmap": HeatMap,
            "wordcloud": WordCloud,
            "boxplot": Boxplot,
            "candlestick": Candlestick,
            "parallel": Parallel,
            "effectScatter": Scatter,
            "graph": Graph,
            "tree": Tree,
            "sankey": Sankey,
            "chord": Chord,
            "liquidFill": Liquid,
            "calendar": HeatMap,  # Calendar uses HeatMap internally
            "polar": Bar,  # Polar uses bar/line with polar coordinate system
            "pictorialBar": PictorialBar,
            "themeRiver": ThemeRiver,
            "bar3D": Bar3D,
            "line3D": Line3D,
            "scatter3D": Scatter3D,
            "surface": Surface3D,
            "lines3D": Lines3D,
            "grid": Grid,
            "treemap": Tree,  # Treemap uses Tree internally
            "sunburst": Tree,  # Sunburst uses Tree internally
            "custom": Bar,  # Custom chart fallback to bar
        }

        ChartClass = chart_map.get(chart_type)
        if ChartClass is None:
            return None

        chart = ChartClass()

        # Extract data based on chart type
        _apply_config_to_chart(chart, config)

        # Render to PNG bytes
        import base64
        png_bytes = chart.render_embed()
        # render_embed returns HTML with embedded SVG/PNG — extract the image
        # For most chart types, pyecharts renders as SVG in render_embed.
        # We need actual PNG. Use render to a temp file instead.
        return _chart_to_png(chart)

    except Exception:
        return None


def _chart_to_png(chart) -> bytes | None:
    """Render a pyecharts chart instance to PNG bytes."""
    try:
        import tempfile
        import subprocess

        with tempfile.NamedTemporaryFile(suffix=".html", delete=False) as f:
            chart.render(f.name)
            html_path = f.name

        png_path = html_path.rstrip(".html") + ".png"
        try:
            # Try using pyecharts built-in screenshot if available
            try:
                from pyecharts.render.snapshot import make_snapshot
                make_snapshot(chart=chart, file_name=png_path)
            except ImportError:
                # Older pyecharts version — try alternate import
                from pyecharts.render import snapshot
                snapshot.make_snapshot(chart=chart, file_name=png_path)
        except Exception:
            # Fallback: render to HTML, then we'll need an external tool
            # For now, return None and let the caller handle gracefully
            os.unlink(html_path)
            return None

        os.unlink(html_path)

        if os.path.exists(png_path):
            with open(png_path, "rb") as f:
                data = f.read()
            os.unlink(png_path)
            return data
        else:
            if os.path.exists(png_path):
                os.unlink(png_path)
            return None
    except Exception:
        return None


def _build_bar_line_effectscatter(chart, series, i, name, data, x_data, config, series_list):
    from pyecharts import options as opts
    if x_data and data and isinstance(data[0], (int, float)):
        chart.add(
            series_name=name,
            x_axis=list(x_data),
            y_axis=list(data),
            is_symbol_show=False,
        )
    elif x_data and data and isinstance(data, list) and len(data) > 0 and isinstance(data[0], dict):
        chart.add(
            series_name=name,
            x_axis=list(x_data),
            y_axis=[d.get("value", d.get("y", 0)) for d in data],
            is_symbol_show=False,
        )
    else:
        chart.add(
            series_name=name,
            x_axis=list(x_data),
            y_axis=list(data),
            is_symbol_show=False,
        )


def _build_pie(chart, series, i, name, data, x_data, config, series_list):
    from pyecharts import options as opts
    if data and isinstance(data[0], dict):
        pie_data = [opts.DataItem(name=d["name"], value=d["value"]) for d in data]
    elif data:
        pie_data = [opts.DataItem(name=str(d), value=d) for d in data]
    else:
        pie_data = []
    chart.add(
        series_name=name,
        data_label_opts=opts.LabelOpts(show=True),
        radius=["40%", "75%"],
        data=pie_data,
    )


def _build_scatter(chart, series, i, name, data, x_data, config, series_list):
    if data and isinstance(data[0], list):
        scatter_data = [tuple(d) for d in data]
    else:
        scatter_data = list(data)
    chart.add(
        series_name=name,
        xaxis_data=list(x_data) if x_data else None,
        series_data=scatter_data,
    )


def _build_radar(chart, series, i, name, data, x_data, config, series_list):
    from pyecharts import options as opts
    indicator = config.get("radar", {}).get("indicator", [])
    radar_indicators = [
        opts.RadarIndicatorItem(name=item.get("name", ""), max_=item.get("max", 100))
        for item in indicator
    ]
    if radar_indicators:
        chart.add_schema(
            schema=radar_indicators,
            start_angle=90,
        )
    chart.add(
        series_name=name,
        data=list(data) if data else [],
    )


def _build_gauge(chart, series, i, name, data, x_data, config, series_list):
    from pyecharts import options as opts
    if data:
        gauge_data = [opts.GaugeDataItem(name=d.get("name", ""), value=d.get("value", 0)) for d in data]
    else:
        gauge_data = []
    chart.add(
        series_name="",
        data=gauge_data,
        detail_opts=opts.LabelOpts(formatter="{value}%"),
    )


def _build_funnel(chart, series, i, name, data, x_data, config, series_list):
    from pyecharts import options as opts
    if data and isinstance(data[0], dict):
        funnel_data = [opts.DataItem(name=d["name"], value=d["value"]) for d in data]
    elif data:
        funnel_data = [opts.DataItem(name=str(d[0]) if isinstance(d, (list, tuple)) else str(d), value=d[1] if isinstance(d, (list, tuple)) and len(d) > 1 else d) for d in data]
    else:
        funnel_data = []
    chart.add(
        series_name=name,
        data=funnel_data,
        sort_="descending",
    )


def _build_heatmap(chart, series, i, name, data, x_data, config, series_list):
    chart.add(
        series_name=name,
        data=list(data) if data else [],
        xaxis_data=list(x_data) if x_data else None,
    )


def _build_wordcloud(chart, series, i, name, data, x_data, config, series_list):
    if data and isinstance(data[0], dict):
        word_data = [(d["name"], d["value"]) for d in data]
    elif data:
        word_data = [(str(d), 1) for d in data]
    else:
        word_data = []
    chart.add(
        series_name=name,
        data_range=[0, max(v for _, v in word_data) if word_data else 1],
        word_size_range=["12", "60"],
        data=word_data,
    )


def _build_boxplot(chart, series, i, name, data, x_data, config, series_list):
    chart.add(
        series_name=name,
        x_axis=list(x_data) if x_data else [],
        y_axis=list(data) if data else [],
    )


def _build_candlestick(chart, series, i, name, data, x_data, config, series_list):
    chart.add(
        series_name=name,
        x_axis=list(x_data) if x_data else [],
        y_axis=list(data) if data else [],
    )


def _build_parallel(chart, series, i, name, data, x_data, config, series_list):
    chart.add(
        series_name=name,
        data=list(data) if data else [],
    )


def _build_graph(chart, series, i, name, data, x_data, config, series_list):
    from pyecharts import options as opts
    nodes = series.get("data", [])
    links = config.get("links", [])
    if nodes:
        node_data = [opts.GraphNode(name=n.get("name", f"Node {i}")) for i, n in enumerate(nodes)]
    else:
        node_data = []
    if links:
        link_data = [opts.GraphLink(source=l.get("source", ""), target=l.get("target", "")) for l in links]
    else:
        link_data = []
    chart.add(
        series_name=name,
        nodes=node_data,
        links=link_data,
        layout=series.get("layout", "force"),
        is_roam=True,
        is_label_show=True,
    )


def _build_tree(chart, series, i, name, data, x_data, config, series_list):
    tree_data = series.get("data", [])
    if tree_data:
        chart.add(
            series_name=name,
            data=tree_data,
            layout=series.get("layout", "orthogonal"),
            orient=series.get("orient", "vertical"),
            symbol_size="auto",
            leaves={},
        )
    else:
        chart.add(series_name=name, data=[], layout="orthogonal")


def _build_sankey(chart, series, i, name, data, x_data, config, series_list):
    from pyecharts import options as opts
    nodes = series.get("data", [])
    links = series.get("links", [])
    if nodes:
        node_data = [opts.SankeyNode(name=n.get("name", f"Node {i}")) for i, n in enumerate(nodes)]
    else:
        node_data = []
    if links:
        link_data = [
            opts.SankeyLink(source=l.get("source", ""), target=l.get("target", ""), value=l.get("value", 0))
            for l in links
        ]
    else:
        link_data = []
    chart.add(
        series_name=name,
        nodes=node_data,
        links=link_data,
        layout_iterations=16,
        orient="horizontal",
        level_distance="auto",
    )


def _build_chord(chart, series, i, name, data, x_data, config, series_list):
    from pyecharts import options as opts
    data = series.get("data", [])
    links = series.get("links", [])
    if data:
        chord_data = [opts.ChordData(name=n.get("name", f"Node {i}")) for i, n in enumerate(data)]
    else:
        chord_data = []
    if links:
        link_data = [
            opts.ChordLink(source=l.get("source", ""), target=l.get("target", ""), value=l.get("value", 0))
            for l in links
        ]
    else:
        link_data = []
    chart.add(
        series_name=name,
        data=chord_data,
        links=link_data,
        tooltip_opts=opts.TooltipOpts(trigger="item"),
    )


def _build_liquidfill(chart, series, i, name, data, x_data, config, series_list):
    if data:
        chart.add(
            series_name=name,
            data=data[:3] if len(data) > 3 else data,
            is_liquid_outline_show=True,
            radius="70%",
        )


def _build_calendar(chart, series, i, name, data, x_data, config, series_list):
    from pyecharts import options as opts
    if data:
        chart.add(
            series_name=name,
            type_="heatmap",
            calendar_index=0,
            data=data,
        )
        chart.set_global_opts(
            visualmap_opts=opts.VisualMapOpts(
                min=data[0][1] if data and len(data[0]) > 1 else 0,
                max=data[0][1] if data and len(data[0]) > 1 else 10,
                orient="horizontal",
                pos_bottom="10",
                pos_left="center",
            )
        )


def _build_polar(chart, series, i, name, data, x_data, config, series_list):
    from pyecharts import options as opts
    if x_data and data and isinstance(data[0], (int, float)):
        chart.add(
            series_name=name,
            x_axis=list(x_data),
            y_axis=list(data),
            coordinate_system="polar",
            label_opts=opts.LabelOpts(is_show=True),
        )


def _build_pictorialbar(chart, series, i, name, data, x_data, config, series_list):
    from pyecharts import options as opts
    if x_data and data and isinstance(data[0], (int, float)):
        chart.add(
            series_name=name,
            x_axis=list(x_data),
            y_axis=list(data),
            symbol_size=["auto", "auto"],
            symbol_margin="25%",
            label_opts=opts.LabelOpts(is_show=True),
        )


def _build_themeriver(chart, series, i, name, data, x_data, config, series_list):
    from pyecharts import options as opts
    if data:
        theme_data = []
        for item in data:
            if isinstance(item, (list, tuple)) and len(item) >= 2:
                theme_data.append((item[0], item[1], item[2] if len(item) > 2 else ""))
            elif isinstance(item, dict):
                theme_data.append((item.get("date", item.get("time", "")), item.get("value", 0), item.get("name", "")))
        chart.add(
            series_name=name,
            data=theme_data,
            singleaxis_opts=opts.SingleAxisOpts(type_="time", pos_top="50", pos_bottom="50"),
        )


def _build_3d(chart, series, i, name, data, x_data, config, series_list):
    from pyecharts import options as opts
    if data:
        chart.add(
            series_name=name,
            data=data,
            shading="lambert",
            label_opts=opts.LabelOpts(is_show=False),
        )
        chart.set_global_opts(
            visualmap_opts=opts.VisualMapOpts(
                max=data[0][2] if data and len(data[0]) > 2 else 100,
                orient="horizontal",
                pos_bottom="0",
                pos_left="center",
            )
        )


def _build_surface(chart, series, i, name, data, x_data, config, series_list):
    from pyecharts import options as opts
    if data:
        chart.add(
            series_name=name,
            data=data,
            shading="color",
            label_opts=opts.LabelOpts(is_show=False),
        )


def _build_lines3d(chart, series, i, name, data, x_data, config, series_list):
    from pyecharts import options as opts
    if data:
        chart.add(
            series_name=name,
            data=data,
            lines3D_opts=opts.Lines3DOpts(
                effect_opts=opts.EffectOpts(
                    period=4,
                    symbol_size=6,
                )
            ),
            label_opts=opts.LabelOpts(is_show=False),
        )


def _build_grid(chart, series, i, name, data, x_data, config, series_list):
    for j, grid_series in enumerate(series_list):
        if not isinstance(grid_series, dict):
            continue
        grid_name = grid_series.get("name", f"Series {j+1}")
        grid_type = grid_series.get("type", "bar")
        grid_data = grid_series.get("data", [])
        grid_xaxis_idx = grid_series.get("xAxisIndex", j % 2)
        grid_yaxis_idx = grid_series.get("yAxisIndex", j % 2)
        if grid_type in ("bar", "line"):
            grid_x_data = []
            if grid_xaxis_idx < len(series_list):
                other_series = series_list[grid_xaxis_idx]
                if isinstance(other_series, dict):
                    other_xaxis = config.get("xAxis", {})
                    if isinstance(other_xaxis, dict):
                        grid_x_data = other_xaxis.get("data", [])
            if grid_x_data and grid_data and isinstance(grid_data[0], (int, float)):
                if grid_type == "bar":
                    chart.add(
                        series_name=grid_name,
                        x_axis=grid_x_data,
                        y_axis=grid_data,
                        xaxis_index=grid_xaxis_idx,
                        yaxis_index=grid_yaxis_idx,
                    )
                else:
                    chart.add(
                        series_name=grid_name,
                        x_axis=grid_x_data,
                        y_axis=grid_data,
                        xaxis_index=grid_xaxis_idx,
                        yaxis_index=grid_yaxis_idx,
                        is_smooth=True,
                    )


def _build_treemap_sunburst(chart, series, i, name, data, x_data, config, series_list):
    if data:
        chart.add(
            series_name=name,
            data=data,
            leaf_depth=2,
            levels=[
                {},
                {"itemstyle": {"color": "#5470c6", "color0": "#91cc75"}},
                {"itemstyle": {"color": "#fac858", "color0": "#ee6666"}},
            ],
        )


def _build_custom(chart, series, i, name, data, x_data, config, series_list):
    if x_data and data and isinstance(data[0], (int, float)):
        chart.add(
            series_name=name,
            x_axis=list(x_data),
            y_axis=list(data),
        )


CHART_BUILDERS = {
    "bar": _build_bar_line_effectscatter,
    "line": _build_bar_line_effectscatter,
    "effectScatter": _build_bar_line_effectscatter,
    "pie": _build_pie,
    "scatter": _build_scatter,
    "radar": _build_radar,
    "gauge": _build_gauge,
    "funnel": _build_funnel,
    "heatmap": _build_heatmap,
    "wordcloud": _build_wordcloud,
    "boxplot": _build_boxplot,
    "candlestick": _build_candlestick,
    "parallel": _build_parallel,
    "graph": _build_graph,
    "tree": _build_tree,
    "sankey": _build_sankey,
    "chord": _build_chord,
    "liquidFill": _build_liquidfill,
    "calendar": _build_calendar,
    "polar": _build_polar,
    "pictorialBar": _build_pictorialbar,
    "themeRiver": _build_themeriver,
    "bar3D": _build_3d,
    "line3D": _build_3d,
    "scatter3D": _build_3d,
    "surface": _build_surface,
    "lines3D": _build_lines3d,
    "grid": _build_grid,
    "treemap": _build_treemap_sunburst,
    "sunburst": _build_treemap_sunburst,
    "custom": _build_custom,
}


def _apply_config_to_chart(chart, config: dict):
    """Apply ECharts config dict to a pyecharts chart instance."""
    from pyecharts import options as opts

    # Apply title
    title_cfg = config.get("title", {})
    if title_cfg:
        text = title_cfg.get("text", "")
        if text:
            chart.set_global_opts(title_opts=opts.TitleOpts(title=text))

    # Apply tooltip
    tooltip_cfg = config.get("tooltip", {})
    if tooltip_cfg:
        trigger = tooltip_cfg.get("trigger", "axis")
        if trigger:
            # Get existing global opts and update tooltip
            current_opts = getattr(chart, "_options", {})
            # pyecharts handles tooltip via set_global_opts
            pass  # tooltip is set via set_global_opts below

    # Extract series data
    series_list = config.get("series", [])
    if not isinstance(series_list, list):
        return

    # Get xAxis data for categorical charts
    xaxis_cfg = config.get("xAxis", {})
    x_data = []
    if isinstance(xaxis_cfg, dict):
        x_data = xaxis_cfg.get("data", [])

    # Get y-axis config
    yaxis_cfg = config.get("yAxis", {})
    y_type = "value" if isinstance(yaxis_cfg, dict) else "value"

    # Process each series
    for i, series in enumerate(series_list):
        if not isinstance(series, dict):
            continue

        name = series.get("name", f"Series {i+1}")
        s_type = series.get("type", "bar")
        data = series.get("data", [])

        builder = CHART_BUILDERS.get(s_type)
        if builder:
            builder(chart, series, i, name, data, x_data, config, series_list)


def _decode_data_url(data_url: str) -> bytes | None:
    """Decode a base64 PNG data URL to raw bytes."""
    try:
        # Format: data:image/png;base64,<base64data>
        if "," in data_url:
            _, b64 = data_url.split(",", 1)
        else:
            b64 = data_url
        import base64
        return base64.b64decode(b64)
    except Exception:
        return None


def process_echarts_blocks(
    text: str,
    chart_pngs: list[str] | None = None,
) -> Tuple[str, list[bytes], list[str]]:
    """Find echarts code blocks in text and replace with chart images.

    chart_pngs: list of base64 PNG data URLs from the frontend (rendered
    via ECharts getDataURL). If provided, used in order. If not provided
    or insufficient, the placeholder is left empty (no image).

    Returns:
        (processed_text, png_bytes_list, raw_json_list)
        - processed_text: original text with echarts blocks replaced by placeholders
        - png_bytes_list: list of PNG bytes for each chart (None entries if no image)
        - raw_json_list: list of raw JSON strings from each block
    """
    png_list: list[bytes] = []
    raw_json_list: list[str] = []

    def _replace_block(match: re.Match) -> str:
        json_str = match.group(1).strip()
        raw_json_list.append(json_str)
        idx = len(raw_json_list) - 1
        placeholder = f"\n[ECHARTS_CHART_{idx}]\n"

        png = None
        if chart_pngs and idx < len(chart_pngs):
            png = _decode_data_url(chart_pngs[idx])
        if png:
            png_list.append(png)
        else:
            png_list.append(b"")  # placeholder to keep indices aligned
        return placeholder

    processed = _ECHARTS_BLOCK_RE.sub(_replace_block, text)
    return processed, png_list, raw_json_list


def _insert_png_in_docx(doc, png_bytes: bytes, index: int):
    """Insert a PNG image into a docx document at the current position."""
    from docx.oxml.ns import qn
    from docx.shared import Inches

    buf = io.BytesIO(png_bytes)
    doc.add_picture(buf, width=Inches(5.5))
    # Add a blank paragraph after the image for spacing
    doc.add_paragraph()


def export_to_pdf(answer_text: str, chart_pngs: list[str] | None = None) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
    from reportlab.lib.units import cm

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
    )
    styles = getSampleStyleSheet()
    body_style = styles["BodyText"]
    body_style.fontSize = 11
    body_style.leading = 16

    # Process echarts blocks BEFORE stripping markdown — the regex
    # needs the ```echarts fences intact to match.
    processed, png_list, _ = process_echarts_blocks(answer_text, chart_pngs)
    clean = _strip_markdown(processed)

    elements = []
    for para in clean.split("\n\n"):
        para = para.strip()
        if not para:
            continue

        # Check if this paragraph is a chart placeholder
        chart_match = re.match(r'\[ECHARTS_CHART_(\d+)\]', para)
        if chart_match:
            idx = int(chart_match.group(1))
            if idx < len(png_list) and png_list[idx]:
                img_buf = io.BytesIO(png_list[idx])
                from reportlab.lib.utils import ImageReader
                img_reader = ImageReader(img_buf)
                # Read image dimensions and scale to fit page width
                avail_width = 612 - doc.leftMargin - doc.rightMargin  # A4 width in points
                orig_w, orig_h = img_reader.getSize()
                scale = min(avail_width / orig_w, 400 / orig_h) if orig_w > 0 else 0.8
                new_w = int(orig_w * scale)
                new_h = int(orig_h * scale)
                # Reset buffer position for Image flowable
                img_buf.seek(0)
                elements.append(Image(img_buf, width=new_w, height=new_h))
                elements.append(Spacer(1, 0.3 * cm))
            continue

        elements.append(Paragraph(para.replace("\n", "<br/>"), body_style))
        elements.append(Spacer(1, 0.3 * cm))

    doc.build(elements)
    return buf.getvalue()


def export_to_word(answer_text: str, chart_pngs: list[str] | None = None) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm, Inches

    doc = Document()
    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Process echarts blocks BEFORE stripping markdown
    processed, png_list, raw_json_list = process_echarts_blocks(answer_text, chart_pngs)
    clean = _strip_markdown(processed)

    for para in clean.split("\n\n"):
        para = para.strip()
        if not para:
            continue

        # Check if this paragraph is a chart placeholder
        chart_match = re.match(r'\[ECHARTS_CHART_(\d+)\]', para)
        if chart_match:
            idx = int(chart_match.group(1))
            # Insert the chart image
            if idx < len(png_list) and png_list[idx]:
                _insert_png_in_docx(doc, png_list[idx], idx)
            # Also include the raw JSON for reference
            if idx < len(raw_json_list):
                json_para = doc.add_paragraph()
                run = json_para.add_run(raw_json_list[idx])
                run.font.size = Pt(8)
                run.font.name = 'Courier New'
            continue

        p = doc.add_paragraph(para)
        p.style.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_image_lines(
    clean: str,
    png_list: list[bytes],
    width: int,
    padding: int,
    wrap_width: int,
) -> list:
    from PIL import Image

    all_lines: list = []
    for para in clean.split("\n"):
        para = para.strip()
        if not para:
            all_lines.append("")
            continue

        chart_match = re.match(r'\[ECHARTS_CHART_(\d+)\]', para)
        if chart_match:
            idx = int(chart_match.group(1))
            if idx < len(png_list) and png_list[idx]:
                chart_img = Image.open(io.BytesIO(png_list[idx]))
                max_h = 400
                ratio = min(width / chart_img.width, max_h / chart_img.height)
                new_w = int(chart_img.width * ratio)
                new_h = int(chart_img.height * ratio)
                chart_img = chart_img.resize((new_w, new_h), Image.Resampling.LANCZOS)
                all_lines.append(("chart", chart_img, new_h))
            continue

        wrapped = textwrap.wrap(para, width=wrap_width) if para else [""]
        all_lines.extend(wrapped or [""])

    return all_lines


def _calculate_image_height(all_lines: list, padding: int, line_height: int) -> int:
    chart_heights = [item[2] for item in all_lines if isinstance(item, tuple)]
    text_count = sum(1 for item in all_lines if isinstance(item, str))
    return padding * 2 + text_count * line_height + sum(chart_heights) + len(chart_heights) * 20


def _render_image_lines(
    all_lines: list,
    img,
    draw,
    padding: int,
    font,
    line_height: int,
) -> None:
    y = padding
    for item in all_lines:
        if isinstance(item, str):
            if item:
                draw.text((padding, y), item, fill=(30, 30, 30), font=font)
            y += line_height
        elif isinstance(item, tuple) and item[0] == "chart":
            _, chart_img, chart_h = item
            img.paste(chart_img, (padding, y))
            y += chart_h + 20


def export_to_image(answer_text: str, chart_pngs: list[str] | None = None) -> bytes:
    from PIL import Image, ImageDraw, ImageFont

    # Process echarts blocks BEFORE stripping markdown
    processed, png_list, _ = process_echarts_blocks(answer_text, chart_pngs)
    clean = _strip_markdown(processed)

    width = 900
    padding = 40
    font_size = 16
    line_height = font_size + 6

    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", font_size)
    except Exception:
        font = ImageFont.load_default()

    wrap_width = (width - 2 * padding) // (font_size // 2 + 2)
    all_lines = _build_image_lines(clean, png_list, width, padding, wrap_width)

    height = _calculate_image_height(all_lines, padding, line_height)

    img = Image.new("RGB", (width, max(height, 200)), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)

    _render_image_lines(all_lines, img, draw, padding, font, line_height)

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def export_message(
    answer_text: str,
    fmt: ExportFormat,
    chart_pngs: list[str] | None = None,
) -> tuple[bytes, str, str]:
    """
    Returns (content_bytes, media_type, filename).
    """
    if fmt == "pdf":
        data = export_to_pdf(answer_text, chart_pngs)
        return data, "application/pdf", "answer.pdf"
    elif fmt == "word":
        data = export_to_word(answer_text, chart_pngs)
        return data, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "answer.docx"
    elif fmt == "image":
        data = export_to_image(answer_text, chart_pngs)
        return data, "image/png", "answer.png"
    else:
        raise ValueError(f"Unknown export format: {fmt}")


